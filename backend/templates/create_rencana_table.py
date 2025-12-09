#!/usr/bin/env python3
"""
Script to create a new template_rps.docx with proper 16-row table
for Rencana Pembelajaran using individual row placeholders.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import shutil

def set_cell_shading(cell, fill_color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_rencana_pembelajaran_table(doc):
    """Create the Rencana Pembelajaran table with 16 rows"""
    
    # Add heading
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("RENCANA PEMBELAJARAN")
    run.bold = True
    run.font.size = Pt(14)
    
    # Create table: 17 rows (1 header + 16 weeks), 8 columns
    table = doc.add_table(rows=17, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths (approximate)
    widths = [Cm(1), Cm(2.5), Cm(4), Cm(4), Cm(3), Cm(1.5), Cm(2), Cm(1.5)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width
    
    # Header row
    headers = [
        "MG KE",
        "KEMAMPUAN AKHIR TIAP TAHAPAN (SUB-CPMK)", 
        "INDIKATOR",
        "TOPIK & SUB-TOPIK MATERI",
        "METODE PEMBELAJARAN (SKEMA BLENDED LEARNING)",
        "WAKTU (MENIT)",
        "TEKNIK & KRITERIA",
        "BOBOT (%)"
    ]
    
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        # Set header style
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        # Blue background
        set_cell_shading(cell, "00B0F0")
    
    # Data rows (16 weeks)
    for week in range(1, 17):
        row = table.rows[week]
        placeholders = [
            f"{{MINGGU_{week}}}",
            f"{{SUB_CPMK_{week}}}",
            f"{{INDIKATOR_{week}}}",
            f"{{TOPIK_{week}}}",
            f"{{METODE_{week}}}",
            f"{{WAKTU_{week}}}",
            f"{{KRITERIA_{week}}}",
            f"{{BOBOT_{week}}}"
        ]
        
        for i, placeholder in enumerate(placeholders):
            cell = row.cells[i]
            cell.text = placeholder
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    return table

def update_existing_template():
    """Update the existing template by replacing the Rencana Pembelajaran section"""
    
    template_path = "template_rps.docx"
    output_path = "template_rps_new.docx"
    
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return
    
    # Load existing template
    doc = Document(template_path)
    
    # Find and note the position of "RENCANA PEMBELAJARAN" table
    # This is complex with python-docx, so we'll create a separate file
    # that can be merged manually
    
    # Create a new document with just the table
    new_doc = Document()
    create_rencana_pembelajaran_table(new_doc)
    new_doc.save("rencana_pembelajaran_table.docx")
    
    print("=" * 60)
    print("CREATED: rencana_pembelajaran_table.docx")
    print("=" * 60)
    print("""
File 'rencana_pembelajaran_table.docx' telah dibuat dengan tabel
yang memiliki 16 row terpisah.

CARA UPDATE TEMPLATE:
1. Buka template_rps.docx di Microsoft Word
2. Hapus tabel 'RENCANA PEMBELAJARAN' yang lama
3. Buka rencana_pembelajaran_table.docx
4. Copy tabel dari rencana_pembelajaran_table.docx
5. Paste ke posisi yang sama di template_rps.docx
6. Simpan template_rps.docx

Placeholder yang digunakan:
- {MINGGU_1} sampai {MINGGU_16}
- {SUB_CPMK_1} sampai {SUB_CPMK_16}
- {INDIKATOR_1} sampai {INDIKATOR_16}
- {TOPIK_1} sampai {TOPIK_16}
- {METODE_1} sampai {METODE_16}
- {WAKTU_1} sampai {WAKTU_16}
- {KRITERIA_1} sampai {KRITERIA_16}
- {BOBOT_1} sampai {BOBOT_16}
    """)

if __name__ == "__main__":
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    update_existing_template()
